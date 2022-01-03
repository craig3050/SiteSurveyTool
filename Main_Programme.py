# General requirements
import requests
import json
import secrets
import datetime
import os
import sys

# QT5 elements
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QApplication, QWidget, QInputDialog, QLineEdit, QFileDialog, QMainWindow
from PyQt5.QtGui import QIcon
from PyQt5.QtGui import QDesktopServices
from PyQt5.QtCore import QUrl

# Word doc elements
from docx import Document
from PIL import Image, ImageOps

# Excel doc elements
from openpyxl import load_workbook

# Logging elements
from exception_decor import exception
from exception_logger import logger

# Airtable elements
airtable_api_key = secrets.airtable_api_key
base_key = secrets.base_key
table_name = secrets.table_name
from pyairtable import Api, Base, Table

from MainUI import Ui_MainWindow

## Some bits for Pyinstaller to help it find relative path files ##
if getattr(sys, 'frozen', False):
    # If the application is run as a bundle, the PyInstaller bootloader
    # extends the sys module by a flag frozen=True and sets the app
    # path into variable _MEIPASS'.
    application_path = sys._MEIPASS
else:
    application_path = os.path.dirname(os.path.abspath(__file__))



class MainWindow:
    def __init__(self):
        self.main_win = QMainWindow()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self.main_win)
        QtCore.QCoreApplication.processEvents()

        self.ui.pushButton_generate_reports.clicked.connect(self.generate_reports)

    def show(self):
        self.main_win.show()
        self.ui.listWidget_outputWindow.clear()
        self.ui.listWidget_outputWindow.addItem("Welcome to the Site Visit Report tool")
        self.ui.listWidget_outputWindow.addItem("Enter the required information then click 'Generate Reports'")
        self.ui.listWidget_outputWindow.addItem("All items must be filled in!")

    def airtable_download(self):
        QtCore.QCoreApplication.processEvents()
        airtable_link = "https://api.airtable.com/v0"
        your_api_key = f"Bearer {airtable_api_key}"
        offset = '0'
        try:
            with requests.session() as response:
                headers = {
                    'Authorization': your_api_key,
                }

                params = {
                    'view':'Grid view',
                    'offset':offset}

                record_list = []
                run = True
                while run is True:
                    response = requests.get(f'{airtable_link}/{base_key}/{table_name}',
                                            headers=headers, params=params)
                    # print(str(response))
                    response = json.loads(response.content)
                    # print(str(response))
                    for item in response['records']:
                        record_list.append(json.dumps(item))
                    if 'offset' in response:
                        offset = str(response['offset'])
                        params = (('offset', offset),)
                    else:
                        run = False
                    print(offset)

                with open('json_output3.json', 'w') as f:
                    f.write(str(record_list))

                return record_list

        except Exception as e:
            print(e)


    def format_airtable_results(self, records):
        QtCore.QCoreApplication.processEvents()
        # Create new dictionary containing all relevant information from Airtable
        entry_dict = {}

        entry_dict['record_id'] = records['id']
        try:
            area = records["fields"]["Area"]
            entry_dict['area'] = area[0]
        except:
            entry_dict['area'] = "Not recorded"
        try:
            entry_dict['observation_number'] = records["fields"]["Observation Number"]
        except:
            entry_dict['observation_number'] = "Not recorded"
        try:
            observation_type = records["fields"]["Observation Type"]
            entry_dict['observation_type'] = observation_type[0]
        except:
            entry_dict['observation_type'] = "Not recorded"
        try:
            entry_dict['description'] = records["fields"]["Description of observation"]
        except:
            entry_dict['description'] = "Not recorded"
        try:
            entry_dict['created_by'] = records["fields"]["Created By"]
        except:
            entry_dict['created_by'] = "Not recorded"
        try:
            print(type(records["fields"]["Created time"]))
            date_object = datetime.datetime.strptime(records["fields"]["Created time"], '%Y-%m-%dT%H:%M:%S.%fZ')
            entry_dict['date'] = date_object.strftime("%d.%m.%Y")
        except Exception as e:
            print(e)
            entry_dict['date'] = "Not recorded"
        try:
            entry_dict['image_link'] = records["fields"]["Attachments"][0]["url"]
        except:
            entry_dict['image_link'] = "Not recorded"
        try:
            entry_dict['status'] = records["fields"]["Status"]
        except:
            entry_dict['status'] = "Open"
        try:
            entry_dict['observation_category'] = records["fields"]["Observation Category"]
        except:
            entry_dict['observation_category'] = "Not recorded"
        try:
            entry_dict['location'] = records["fields"]["Location"]
        except:
            entry_dict['location'] = "Not recorded"
        return entry_dict


    def download_picture(self, picture_link, observation_number):
        directory_name = 'Pictures'
        try:
            try:
                os.makedirs(directory_name)
            except:
                print("Directory already exists")

            if os.path.isfile(f'{directory_name}/{observation_number}.jpg'):
                print("File already exists")
            else:
                print("File does not exist")
                img_data = requests.get(picture_link).content
                try:
                    with open(f'{directory_name}/{observation_number}.jpg', 'wb') as handler:
                        handler.write(img_data)
                except Exception as e:
                    print(e)
                try:
                    image_pil = Image.open(f'{directory_name}/{observation_number}.jpg').convert('RGB')
                    image_pil = ImageOps.exif_transpose(image_pil)
                    image_pil.save(f'{directory_name}/{observation_number}.jpg')
                except:
                    "Cannot transpose image"
        except:
            print('No image to download')

        return f'{directory_name}/{observation_number}.jpg'


    def export_to_excel(self, formatted_results, image_link, svr_no):
        # if workbook is available
        # open workbook
        # else
        # create new workbook from template
        if os.path.isfile('The Factory - Site Observation Reports Summary.xlsx'):
            # print("File already exists")
            wb = load_workbook("The Factory - Site Observation Reports Summary.xlsx")
        else:
            wb = load_workbook("Summary Template.xlsx")

        ws = wb['Sheet1']

        # check if entry is in current column - exit if true
        temp_list = []
        for cell in ws['A']:
            temp_list.append(cell.value)
        print(temp_list)

        if str(formatted_results['record_id']) not in temp_list:
            print('Excel Sheet - Record not present')
            today = datetime.date.today()
            today = today.strftime("%d/%m/%Y")
            ws.cell(column=2, row=4, value=str(today))

            # if entry is not in column
            new_row_location = ws.max_row + 1
            # write to the cell you want, specifying row and column, and value
            ws.cell(column=1, row=new_row_location, value=str(formatted_results['record_id']))
            ws.cell(column=2, row=new_row_location, value=str(formatted_results['observation_number']))
            ws.cell(column=3, row=new_row_location, value=str(svr_no))
            ws.cell(column=5, row=new_row_location, value=str(formatted_results['date']))
            ws.cell(column=6, row=new_row_location, value=str(formatted_results['area']))
            ws.cell(column=7, row=new_row_location, value=str(formatted_results['location']))
            ws.cell(column=8, row=new_row_location, value=str(formatted_results['description']))
            ws.cell(column=9, row=new_row_location, value=str(formatted_results['status']))
            ws.cell(column=10, row=new_row_location, value=str(formatted_results['observation_type']))
            ws.cell(column=11, row=new_row_location, value=str(formatted_results['observation_category']))
            hyperlink_location = f'=HYPERLINK(O{new_row_location})'
            ws.cell(column=12, row=new_row_location, value=hyperlink_location)
            ws.cell(column=14, row=new_row_location, value='=LEFT(CELL("filename", A1), FIND("[",CELL("filename",A1))-1)')
            ws.cell(column=15, row=new_row_location, value=str(image_link))
        else:
            print('Excel Sheet - Record already present')


        # Save the sheet and close
        wb.save(filename='The Factory - Site Observation Reports Summary.xlsx')
        wb.close()


    @exception(logger)
    def export_to_word(self, formatted_results, image_link, svr_no, site_visit_date, surveyor_names, chaperone,
                       issued_by, progress_from_site):

        if os.path.isfile(f'SiteReport.docx'):
            # print("File already exists")
            doc = Document("SiteReport.docx")
        else:
            doc = Document("Template.docx")

        QtCore.QCoreApplication.processEvents()

        # Add site report number (take number from input)
        doc.tables[0].cell(0, 1).text = svr_no

        # Add today's date to the table
        today = datetime.date.today()
        today = today.strftime("%d/%m/%Y")
        doc.tables[0].cell(2, 1).text = str(today)

        # Add name to 'issued by' field (take from input)
        doc.tables[1].cell(1, 1).text = issued_by

        # Replace placeholder text with the relevant info (take from input)
        placeholder1 = f"""{surveyor_names} (BDP) attended site {site_visit_date} for a design team meeting and general site walkover with the construction team.

{chaperone} from the Contractor team was in attendance.

Progress on site includes:
{progress_from_site}
"""
        for text_to_replace in doc.paragraphs:
            if "<<Placeholder1>>" in text_to_replace.text:
                text_to_replace.text = placeholder1

        # Add information from survey into the table
        row_count = len(doc.tables[2].rows)
        doc.tables[2].cell(row_count - 1, 0).text = str(formatted_results['observation_number'])
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("Area:\n").bold = True
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['area']))
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nLocation:\n").bold = True
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['location']))
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nObservation Type:\n").bold = True
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['observation_type']))
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nObservation Category:\n").bold = True
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['observation_category']))
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nDescription:\n").bold = True
        doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['description']))

        # Add empty paragraph in the photos column so the photo can be added
        paragraph = doc.tables[2].cell(row_count - 1, 2).paragraphs[0]

        try:
            run = paragraph.add_run()
            run.add_picture(image_link, width=3000000)
        except Exception as e:
            print(e)
        doc.tables[2].add_row()  # ADD ROW HERE
        doc.save("SiteReport.docx")


    def generate_reports(self):
        svr_no = self.ui.lineEdit_svr_no.text()
        site_visit_date = self.ui.lineEdit_site_visit_date.text()
        date1 = self.ui.lineEdit_date1.text()
        date2 = self.ui.lineEdit_date2.text()
        surveyor_names = self.ui.lineEdit_present_for_survey.text()
        chaperone = self.ui.lineEdit_chaperone.text()
        issued_by = self.ui.lineEdit_issued_by.text()
        progress_from_site = self.ui.textEdit_progress_notes.toPlainText()

        if svr_no == "":
            self.ui.listWidget_outputWindow.addItem("SVR number empty - please add value")
            return 0
        if site_visit_date == "":
            self.ui.listWidget_outputWindow.addItem("Site Visit Date empty - please add value")
            return 0
        if date1 == "":
            self.ui.listWidget_outputWindow.addItem("Start date empty - please add value")
            return 0
        if date2 == "":
            self.ui.listWidget_outputWindow.addItem("End date empty - please add value")
            return 0
        if surveyor_names == "":
            self.ui.listWidget_outputWindow.addItem("Surveyor Names empty - please add value")
            return 0
        if chaperone == "":
            self.ui.listWidget_outputWindow.addItem("Chaperone name empty - please add value")
            return 0
        if issued_by == "":
            self.ui.listWidget_outputWindow.addItem("Issued by field empty - please add value")
            return 0
        if progress_from_site == "":
            self.ui.listWidget_outputWindow.addItem("Please add progress notes from site")
            return 0
        else:
            self.ui.listWidget_outputWindow.clear()
            self.ui.listWidget_outputWindow.addItem("Carrying on")

        # Download information from Airtable
        airtable_response = self.airtable_download()
        # print(json.dumps(airtable_response, indent=4))

        date1 = datetime.datetime.strptime(date1, "%d/%m/%Y")
        date2 = datetime.datetime.strptime(date2, "%d/%m/%Y")

        try:
            self.ui.listWidget_outputWindow.clear()
            # Send download data format into correct sections
            # for records in airtable_response['records']:
            for records in airtable_response:
                QtCore.QCoreApplication.processEvents()
                print(records)
                records = json.loads(records)
                for item in records:
                    print(item)
                # Download all the items from Airtable
                formatted_results = self.format_airtable_results(records)

                # Download all the new images - this will skip existing images
                image_link = self.download_picture(formatted_results['image_link'], formatted_results['observation_number'])

                print_observation_number = formatted_results['observation_number']
                # Add the items to a Word report
                # Look through the returned results and only incorporate in date range
                self.ui.listWidget_outputWindow.addItem(f'Processing {print_observation_number}')

                survey_date = datetime.datetime.strptime(formatted_results['date'], "%d.%m.%Y")
                # Check if survey was between the two dates specified
                if date1 <= survey_date <= date2:
                    print(f'Moving on to export to word using date {survey_date}')
                    self.export_to_word(formatted_results, image_link, svr_no, site_visit_date, surveyor_names,
                                        chaperone, issued_by, progress_from_site)

                # Add the items to a master Excel document
                print("Now exporting to Excel")
                self.export_to_excel(formatted_results, image_link, svr_no)

            self.ui.listWidget_outputWindow.clear()
            self.ui.listWidget_outputWindow.addItem("Report generation complete")


        except Exception as e:
            print(e)

if __name__ == '__main__':
    QtWidgets.QApplication.setAttribute(QtCore.Qt.ApplicationAttribute.AA_EnableHighDpiScaling, True)
    QtWidgets.QApplication.setAttribute(QtCore.Qt.ApplicationAttribute.AA_UseHighDpiPixmaps, True)
    app = QApplication(sys.argv)
    main_win = MainWindow()
    main_win.show()
    sys.exit(app.exec_())
