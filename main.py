import requests
import json
import secrets
import datetime
import os
from exception_decor import exception
from exception_logger import logger
from docx import Document
from PIL import Image, ImageOps
from openpyxl import load_workbook

airtable_api_key = secrets.airtable_api_key
base_key = secrets.base_key
table_name = secrets.table_name

@exception(logger)
def airtable_download():
    airtable_link = "https://api.airtable.com/v0"
    your_api_key = f"Bearer {airtable_api_key}"

    try:
        with requests.session() as response:
            headers = {
                'Authorization': your_api_key,
            }

            params = (
                ('maxRecords', '1000'),
                ('view', 'Grid view'),
            )
            response = requests.get(f'{airtable_link}/{base_key}/{table_name}',
                                    headers=headers, params=params)

            #print(str(response))
            #print(str(response.content))
            json_to_print = json.loads(response.content)
            # print(json.dumps(json_to_print, indent=4))

            return json_to_print

    except Exception as e:
        print(e)

@exception(logger)
def format_airtable_results(records):
    # Create new dictionary containing all relevant information from Airtable
    entry_dict = {}

    entry_dict['record_id'] = records['id']
    try:
        area = records["fields"]["Area"]
        entry_dict['area'] = area[0]
    except:
        entry_dict['area'] = "Not recorded"
    try:
        entry_dict['observation_number'] = records["fields"]["Observation Number"]
    except:
        entry_dict['observation_number'] = "Not recorded"
    try:
        observation_type = records["fields"]["Observation Type"]
        entry_dict['observation_type'] = observation_type[0]
    except:
        entry_dict['observation_type'] = "Not recorded"
    try:
        entry_dict['description'] = records["fields"]["Description of observation"]
    except:
        entry_dict['description'] = "Not recorded"
    try:
        entry_dict['created_by'] = records["fields"]["Created By"]
    except:
        entry_dict['created_by'] = "Not recorded"
    try:
        date_object = datetime.datetime.strptime(records["fields"]["Created time"], '%Y-%m-%dT%H:%M:%S.%fZ')
        entry_dict['date'] = date_object.strftime("%d.%m.%Y")
    except Exception as e:
        print(e)
        entry_dict['date'] = "Not recorded"
    try:
        entry_dict['image_link'] = records["fields"]["Attachments"][0]["url"]
    except:
        entry_dict['image_link'] = "Not recorded"
    try:
        entry_dict['status'] = records["fields"]["Status"]
    except:
        entry_dict['status'] = "Open"
    try:
        entry_dict['observation_category'] = records["fields"]["Observation Category"]
    except:
        entry_dict['observation_category'] = "Not recorded"
    try:
        entry_dict['location'] = records["fields"]["Location"]
    except:
        entry_dict['location'] = "Not recorded"
    return entry_dict


def export_to_excel(formatted_results, image_link, svr_no):
    #if workbook is available
    #open workbook
    #else
    #create new workbook from template
    if os.path.isfile('The Factory - Site Observation Reports Summary.xlsx'):
        #print("File already exists")
        wb = load_workbook("The Factory - Site Observation Reports Summary.xlsx")
    else:
        wb = load_workbook("Summary Template.xlsx")

    ws = wb['Sheet1']

    #check if entry is in current column - exit if true
    for cell in ws['A']:
        if cell == str(formatted_results['record_id']):
            print('Record already present')
            return 0

    today = datetime.date.today()
    today = today.strftime("%d/%m/%Y")
    ws.cell(column=2, row=4, value=str(today))

    #if entry is not in column
    new_row_location = ws.max_row + 1
    # write to the cell you want, specifying row and column, and value
    ws.cell(column=1, row=new_row_location, value=str(formatted_results['record_id']))
    ws.cell(column=2, row=new_row_location, value=str(formatted_results['observation_number']))
    ws.cell(column=3, row=new_row_location, value=str(svr_no))
    ws.cell(column=5, row=new_row_location, value=str(formatted_results['date']))
    ws.cell(column=6, row=new_row_location, value=str(formatted_results['area']))
    ws.cell(column=7, row=new_row_location, value=str(formatted_results['location']))
    ws.cell(column=8, row=new_row_location, value=str(formatted_results['description']))
    ws.cell(column=9, row=new_row_location, value=str(formatted_results['status']))
    ws.cell(column=10, row=new_row_location, value=str(formatted_results['observation_type']))
    ws.cell(column=11, row=new_row_location, value=str(formatted_results['observation_category']))
    hyperlink_location = f'=HYPERLINK(O{new_row_location})'
    ws.cell(column=12, row=new_row_location, value=hyperlink_location)
    ws.cell(column=14, row=new_row_location, value='=LEFT(CELL("filename", A1), FIND("[",CELL("filename",A1))-1)')
    ws.cell(column=15, row=new_row_location, value=str(image_link))

    # Save the sheet and close
    wb.save(filename='The Factory - Site Observation Reports Summary.xlsx')
    wb.close()

@exception(logger)
def export_to_word(formatted_results, image_link, svr_no):
    if os.path.isfile(f'SiteReport.docx'):
        #print("File already exists")
        doc = Document("SiteReport.docx")
    else:
        doc = Document("Template.docx")

    # Add site report number (take number from input)
    doc.tables[0].cell(0, 1).text = svr_no

    # Add today's date to the table
    today = datetime.date.today()
    today = today.strftime("%d/%m/%Y")
    doc.tables[0].cell(2, 1).text = str(today)

    # Add name to 'issued by' field (take from input)
    issued_by = "Craig Cuninghame"
    doc.tables[1].cell(1, 1).text = issued_by

    # Replace placeholder text with the relevant info (take from input)
    surveyor_names = "Craig Cuninghame & Gary Fairclough"
    date_of_survey = "<<date from input>>"
    chaperone = "<<Person who walked us round>>"
    progress_from_site = "Nothing so far..."
    placeholder1 = f"""{surveyor_names} (BDP) attended site {date_of_survey} for a design team meeting and general site walkover with the construction team.

{chaperone} from the Contractor team was in attendance.

Progress on site includes:
{progress_from_site}
"""
    for text_to_replace in doc.paragraphs:
        if "<<Placeholder1>>" in text_to_replace.text:
            text_to_replace.text = placeholder1

    # Add information from survey into the table
    row_count = len(doc.tables[2].rows)
    doc.tables[2].cell(row_count - 1, 0).text = str(formatted_results['observation_number'])
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("Area:\n").bold=True
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['area']))
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nLocation:\n").bold=True
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['location']))
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nObservation Type:\n").bold=True
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['observation_type']))
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nObservation Category:\n").bold=True
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['observation_category']))
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run("\nDescription:\n").bold=True
    doc.tables[2].cell(row_count - 1, 1).paragraphs[0].add_run(str(formatted_results['description']))

    # Add empty paragraph in the photos column so the photo can be added
    paragraph = doc.tables[2].cell(row_count-1, 2).paragraphs[0]

    try:
        run = paragraph.add_run()
        run.add_picture(image_link, width = 3000000)
    except Exception as e:
        print(e)
    doc.tables[2].add_row() #ADD ROW HERE
    doc.save("SiteReport.docx")

@exception(logger)
def download_picture(picture_link, observation_number):
    directory_name = 'Pictures'
    try:
        try:
            os.makedirs(directory_name)
        except:
            print("Directory already exists")

        if os.path.isfile(f'{directory_name}/{observation_number}.jpg'):
            print("File already exists")
        else:
            print("File does not exist")
            img_data = requests.get(picture_link).content
            try:
                with open(f'{directory_name}/{observation_number}.jpg', 'wb') as handler:
                    handler.write(img_data)
            except Exception as e:
                print(e)
            try:
                image_pil = Image.open(f'{directory_name}/{observation_number}.jpg').convert('RGB')
                image_pil = ImageOps.exif_transpose(image_pil)
                image_pil.save(f'{directory_name}/{observation_number}.jpg')
            except:
                "Cannot transpose image"
    except:
        print('No image to download')

    return f'{directory_name}/{observation_number}.jpg'



if __name__ == '__main__':
    svr_no = "001"

    # Download information from Airtable
    airtable_response = airtable_download()
    #print(json.dumps(airtable_response, indent=4))

    # Send download data format into correct sections
    for records in airtable_response['records']:
        # Download all the items from Airtable
        formatted_results = format_airtable_results(records)

        # Download all the new images - this will skip existing images
        image_link = download_picture(formatted_results['image_link'], formatted_results['observation_number'])

        # Add the items to a Word report
        # Look through the returned results and only incorporate in date range
        date1 = "20/01/2021" #change to input
        date2 = "21/05/2021" #change to input
        date1 = datetime.datetime.strptime(date1, "%d/%m/%Y")
        date2 = datetime.datetime.strptime(date2, "%d/%m/%Y")
        survey_date = datetime.datetime.strptime(formatted_results['date'], "%d.%m.%Y")
        # Check if survey was between the two dates specified
        if date1 <= survey_date <= date2:
            export_to_word(formatted_results, image_link, svr_no)

        # Add the items to a master Excel document
        export_to_excel(formatted_results, image_link, svr_no)


